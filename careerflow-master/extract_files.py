import os
from docx import Document

# Change this to your project folder path
project_folder = r"C:\Users\APOORVA\Desktop\innov\careerflow-master"

# Create a Word Document
doc = Document()
doc.add_heading('Project Files Content', level=1)

# Allowed text-based file extensions
TEXT_FILE_EXTENSIONS = ('.txt', '.md', '.js', '.py', '.html', '.css', '.json', '.ts', '.jsx', '.tsx', '.env')

# Function to check if a file is text-based
def is_text_file(file_path):
    return file_path.endswith(TEXT_FILE_EXTENSIONS)

# Function to add file content to the document
def add_file_to_doc(file_path):
    try:
        if is_text_file(file_path):
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

                # If content is too long, truncate it
                if len(content) > 5000:  # Adjust limit if needed
                    content = content[:5000] + "\n... [Content Truncated]"

                doc.add_paragraph(f"📂 File Path: {file_path}\n", style="Heading2")
                doc.add_paragraph(content.encode('utf-8', 'ignore').decode('utf-8'))  # Ensure UTF-8 encoding
                doc.add_paragraph("\n" + "="*50 + "\n")  # Separator
    except Exception as e:
        print(f"❌ Error reading {file_path}: {e}")

# Walk through all files in the project folder
for root, _, files in os.walk(project_folder):
    for file in files:
        file_path = os.path.join(root, file)
        add_file_to_doc(file_path)

# Save the document
output_file = os.path.join(project_folder, "Project_Files_Content.docx")
doc.save(output_file)
print(f"✅ All files saved into '{output_file}'")
